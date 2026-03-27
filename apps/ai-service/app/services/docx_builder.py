import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ACCENT_COLOR = RGBColor(0x1F, 0x2A, 0x44)
BLACK = RGBColor(0x00, 0x00, 0x00)
FONT_NAME = "Calibri"

METRIC_PATTERN = re.compile(r"(\$?\d[\d,]*(?:\.\d+)?%?|\d+\+|\d+(?:\.\d+)?[KMB]?)")
TECH_PATTERN = re.compile(
    r"\b(React|TypeScript|JavaScript|Node\.js|Node|Python|PostgreSQL|SQL|GraphQL|REST APIs|Docker|AWS|Azure|GCP|Kubernetes|Terraform|Next\.js|Tailwind|Redis|CI/CD)\b",
    re.IGNORECASE,
)


def build_docx(file_path: str, title: str, content: str) -> str:
    path = Path(file_path)
    path.parent.mkdir(parents=True, exist_ok=True)

    document = Document()
    _configure_document(document)

    lines = [line.rstrip() for line in content.splitlines()]
    non_empty_lines = [line.strip() for line in lines if line.strip()]
    if not non_empty_lines:
        non_empty_lines = [title]

    header_lines = _extract_header_lines(non_empty_lines, title)
    _add_header(document, header_lines, title)

    remaining_lines = _trim_used_lines(lines, header_lines)
    _render_body(document, remaining_lines)

    document.save(path)
    return str(path)


def _configure_document(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    style = document.styles["Normal"]
    style.font.name = FONT_NAME
    style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    style.font.size = Pt(11)


def _extract_header_lines(lines: list[str], fallback_title: str) -> list[str]:
    if len(lines) >= 4 and "@" not in lines[0]:
        return lines[:4]
    return [
        "Jose Puentes",
        fallback_title,
        "Miami, FL | +1 (339) 400 8238",
        "josepuentes0207@gmail.com | LinkedIn",
    ]


def _trim_used_lines(lines: list[str], header_lines: list[str]) -> list[str]:
    trimmed = list(lines)
    for header_line in header_lines:
        for index, line in enumerate(trimmed):
            if line.strip() == header_line.strip():
                trimmed = trimmed[index + 1 :]
                break
    return trimmed


def _add_header(document: Document, header_lines: list[str], fallback_title: str) -> None:
    name_line = header_lines[0] if header_lines else "Jose Puentes"
    role_line = header_lines[1] if len(header_lines) > 1 else fallback_title
    detail_line = header_lines[2] if len(header_lines) > 2 else "Miami, FL | +1 (339) 400 8238"
    contact_line = (
        header_lines[3] if len(header_lines) > 3 else "josepuentes0207@gmail.com | LinkedIn"
    )

    name = document.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name.space_after = Pt(2)
    run = name.add_run(name_line)
    _style_run(run, 25, True, ACCENT_COLOR)

    role = document.add_paragraph()
    role.alignment = WD_ALIGN_PARAGRAPH.CENTER
    role.space_after = Pt(2)
    run = role.add_run(role_line)
    _style_run(run, 16, True, BLACK)

    details = document.add_paragraph()
    details.alignment = WD_ALIGN_PARAGRAPH.CENTER
    details.space_after = Pt(0)
    run = details.add_run(detail_line)
    _style_run(run, 11, False, BLACK)

    contact = document.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.space_after = Pt(14)
    run = contact.add_run(contact_line)
    _style_run(run, 11, False, BLACK)


def _render_body(document: Document, lines: list[str]) -> None:
    index = 0
    while index < len(lines):
        line = lines[index].strip()
        if not line:
            index += 1
            continue

        if _is_section_header(line):
            _add_section_header(document, line)
            index += 1
            continue

        if line.startswith("- "):
            _add_bullet(document, line[2:].strip())
            index += 1
            continue

        next_line = lines[index + 1].strip() if index + 1 < len(lines) else ""
        if " | " in next_line and not next_line.startswith("- ") and not _is_section_header(next_line):
            _add_role_header(document, line, next_line)
            index += 2
            continue

        _add_body_paragraph(document, line)
        index += 1


def _is_section_header(line: str) -> bool:
    normalized = line.replace(" ", "").replace("&", "")
    return normalized.isupper() and len(line) > 3


def _add_section_header(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.space_before = Pt(12)
    paragraph.space_after = Pt(6)
    run = paragraph.add_run(text.upper())
    _style_run(run, 14, True, ACCENT_COLOR)


def _add_role_header(document: Document, title_line: str, meta_line: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.space_before = Pt(10)
    paragraph.space_after = Pt(0)
    run = paragraph.add_run(_remove_scope_text(title_line))
    _style_run(run, 11, True, BLACK)

    meta = document.add_paragraph()
    meta.space_after = Pt(4)
    run = meta.add_run(meta_line)
    _style_run(run, 11, False, BLACK)
    run.italic = True


def _add_body_paragraph(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.space_after = Pt(6)
    _write_text_with_emphasis(paragraph, text)


def _add_bullet(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.space_after = Pt(6)
    _write_text_with_emphasis(paragraph, text, is_bullet=True)


def _write_text_with_emphasis(paragraph, text: str, is_bullet: bool = False) -> None:
    fragments = _emphasize_fragments(text, is_bullet=is_bullet)
    if not fragments:
        fragments = [(text, False)]

    for value, is_bold in fragments:
        run = paragraph.add_run(value)
        _style_run(run, 11, is_bold, BLACK)


def _emphasize_fragments(text: str, is_bullet: bool) -> list[tuple[str, bool]]:
    spans: list[tuple[int, int]] = []

    if is_bullet:
        verb_end = len(text.split(" ", 1)[0])
        spans.append((0, verb_end))

    for match in METRIC_PATTERN.finditer(text):
        spans.append(match.span())

    tech_matches = list(TECH_PATTERN.finditer(text))[:2]
    for match in tech_matches:
        spans.append(match.span())

    spans = _merge_spans(spans)
    if not spans:
        return [(text, False)]

    fragments: list[tuple[str, bool]] = []
    cursor = 0
    for start, end in spans:
        if start > cursor:
            fragments.append((text[cursor:start], False))
        fragments.append((text[start:end], True))
        cursor = end
    if cursor < len(text):
        fragments.append((text[cursor:], False))
    return fragments


def _merge_spans(spans: list[tuple[int, int]]) -> list[tuple[int, int]]:
    if not spans:
        return []
    sorted_spans = sorted(spans)
    merged = [sorted_spans[0]]
    for start, end in sorted_spans[1:]:
        last_start, last_end = merged[-1]
        if start <= last_end:
            merged[-1] = (last_start, max(last_end, end))
        else:
            merged.append((start, end))
    return merged


def _remove_scope_text(value: str) -> str:
    return re.sub(r"\s*\([^)]*\)", "", value).strip()


def _style_run(run, size: int, bold: bool, color: RGBColor) -> None:
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
